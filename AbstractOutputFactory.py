# *coding=utf-8
from __future__ import annotations
from abc import ABC,abstractmethod
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

#Abstract factories for output
class AbstractOutputFactory(ABC):
    @abstractmethod
    def createFormDefault(self):
        pass

#Concrete factories which inhere from AbstractOutputFactory
class OutputForWord(AbstractOutputFactory):
    def __init__(self,docPath=None):
        self.docPath = docPath

    def createFormDefault(self): 
        return WordFormDefault(self.docPath)

#Abstract class for word output
class AbstractOutputFormWord(ABC):
    @abstractmethod
    def write(self):
        pass
    @abstractmethod
    def OutputTitle(self):
        pass
    @abstractmethod
    def OutputParagraph(self):
        pass

    @abstractmethod
    def PageBreak(self):
        pass

    @abstractmethod
    def ListStylesName(self):
        pass

    @abstractmethod
    def setStyle(self):
        pass

class WordFormDefault(AbstractOutputFormWord):
    def __init__(self,docPath=None):
        self.document = Document()
        self.savePath = docPath

    def write(self,savePath=None):
        if savePath != None:
            self.document.save(savePath)
        else:
            self.document.save(self.savePath)

    """
    @Usage: create and write title to target
    @level: level of title
    @left_ident: left ident,default set to 0;if center equal 1,this param is not valied
    @center: if set to 1,title will center align
    @bold: default set to 1,which is yes
    @size: default set to 14,just a random number
    @rgbColor: default set to 0x00,0x00,0x00
    @typeFace: default set to '黑体'
    """
    def OutputTitle(self,content=None,level=1,left_ident=0,center=1,bold=1,size=14,rgbColor=None,typeFace=u"黑体"):

        #Default title is "Default heading"
        if content != None:
            p = self.document.add_heading("",level)
            p_run = p.add_run(content)
        else:
            p = self.document.add_heding("",level)
            p_run = p.add_run("Default heading")

        #Set style
        if center == 1:
            self.setStyle(p,p_run,alignment=WD_ALIGN_PARAGRAPH.CENTER,typeFace=typeFace)
        else:
            self.setStyle(p,p_run,left_ident=left_ident,typeFace=typeFace)
        #Write to target file
        self.write(savePath=None)
    """
    @Usage: create and write paragraph to target file
    @content: Content you want write to your file
    @left_ident: left ident,default set to 0;if center equal 1,this param is not valied
    @center: if set to 1,title will center align,default is 0
    @bold: default set to 1,which is yes
    @size: default set to 14,just a random number
    @rgbColor: default set to 0x00,0x00,0x00
    @typeFace: default set to '宋体'
    """
    def OutputParagraph(self,content=None,left_ident=0,
                        center=0,bold=0,size=12,rgbColor=None,
                        typeFace=u"宋体"):
        if content != None:
            p = self.document.add_paragraph("")
            p_run = p.add_run(content)
        else:
            p = self.document.add_paragraph("")
            p_run = p.add_run("Paragraph example")

        #Set style
        if center == 1:
            self.setStyle(p,p_run,center=WD_ALIGN_PARAGRAPH.CENTER,bold=bold,fontSize=size)_
        else:
            self.setStyle(p,p_run,left_ident=left_ident,bold=bold,fontSize=size)

        #Write to target file
        self.write(savePath=None)

    def PageBreak(self):
         self.document.add_page_break()

    def ListStylesName(self):
        styleArray = [
                s for s in self.document.styles 
                ]
        for i in styleArray:
            print(i.name)

    """
    @Usage: Set style for every component
    @so: style object 
    @s_run: style run object
    @alignment: alignment paramter
    @left_ident: left ident,default set to 0;if center equal 1,this param is not valied
    @bold: default set to 1,which is yes
    @fontSize: default set to 14,just a random number
    @rgbColor: default set to 0x00,0x00,0x00
    @typeFace: default set to '宋体'
    """
    def setStyle(self,so,so_run,alignment=None,left_ident=0,
                fontSize=14,bold=1,rgbColor=RGBColor(0x00,0x00,0x00),
                typeFace=u"宋体"):

        #Set paragraph format
        paragraph_format = so.style.paragraph_format
        if alignment != None:
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph_format.left_indent = Inches(left_ident)

        #Set font attributes
        font = so.style.font
        font.bold = bold
        font.size = Pt(fontSize)

        #Default color is black
        font.color.rgb = rgbColor

        #Set font type using Run object
        so_run.font.name = typeFace
        so_run._element.rPr.rFonts.set(qn('w:eastAsia'),typeFace)

